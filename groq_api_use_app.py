from flask import Flask, request, jsonify, send_file, url_for, make_response
import os
import uuid
import io
import re
from groq import Groq
from flask_caching import Cache
from werkzeug.exceptions import BadRequest, NotFound, InternalServerError
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from docx import Document

app = Flask(__name__)

# Set up caching
app.config['CACHE_TYPE'] = 'simple'
cache = Cache(app)

# Set up rate limiting
limiter = Limiter(
    key_func=get_remote_address,
    app=app,
    default_limits=["200 per day", "50 per hour"],
    storage_uri="memory://",
)

# Set the Groq API key (use environment variables in production)
GROQ_API_KEY = "Here I pase my groq website llama3.2 api"
client = Groq(api_key=GROQ_API_KEY)

SYSTEM_MESSAGE_EN = """
You are "KUROCO LAB chatbot", created by JB Connect Ltd. As a managing director and project implementor, your role is to:
1. Guide users through project descriptions, asking relevant questions to gather comprehensive information.
2. Produce detailed project summaries covering: project name, goals, scope, timeline, budget, resources needed, risks, and key stakeholders.
3. Offer industry-specific insights and best practices for project management.
4. Provide approximate cost and timeline estimates based on project complexity.
5. Suggest potential challenges and mitigation strategies.
6. Recommend project management methodologies suitable for the described project.
7. Offer to create and provide downloadable SRS documents when appropriate.
Always maintain a professional yet approachable tone. Be proactive in seeking clarification and offering additional information to ensure comprehensive project planning.
"""

SYSTEM_MESSAGE_JP = """
あなたは、JB Connect Ltd.が作成した「KUROCO LABチャットボット」です。マネージングディレクターおよびプロジェクト実施者として、あなたの役割は以下の通りです：
1. ユーザーにプロジェクトの説明を促し、関連する質問をして包括的な情報を収集する。
2. プロジェクト名、目標、範囲、タイムライン、予算、必要なリソース、リスク、主要な利害関係者を含む詳細なプロジェクト要約を作成する。
3. 業界固有の洞察とプロジェクト管理のベストプラクティスを提供する。
4. プロジェクトの複雑さに基づいて、概算のコストとタイムラインの見積もりを提供する。
5. 潜在的な課題と緩和策を提案する。
6. 説明されたプロジェクトに適したプロジェクト管理手法を推奨する。
7. 適切な場合、ダウンロード可能なSRSドキュメントの作成と提供を申し出る。
常にプロフェッショナルでありながら親しみやすい口調を維持してください。明確化を求め、追加情報を提供することで、包括的なプロジェクト計画を確実にするよう積極的に行動してください。
"""

documents = {}
conversation_history = []
user_language = 'en'  # Default language

def process_assistant_message(assistant_message, user_message):
    global user_language
    if any(keyword in user_message.lower() for keyword in ["document", "report", "summary", "download", "link", "srs"]):
        doc_id = str(uuid.uuid4())
        srs_content = generate_srs_content(conversation_history)
        documents[doc_id] = srs_content
        download_link = url_for('get_document', doc_id=doc_id, _external=True)
        if user_language == 'en':
            assistant_message += f"\n\nI've prepared an SRS document based on our conversation. Here's the link to download your SRS document: [Download SRS Document]({download_link})"
        else:
            assistant_message += f"\n\n会話に基づいてSRSドキュメントを作成しました。以下のリンクからSRSドキュメントをダウンロードできます：[SRSドキュメントをダウンロード]({download_link})"
    return assistant_message

def generate_srs_content(conversation_history):
    conversation_text = "\n".join([f"{'Human' if i % 2 == 0 else 'Assistant'}: {msg}" for i, msg in enumerate(conversation_history)])
    
    srs_prompt = f"""
    Based on the following conversation, generate a comprehensive Software Requirements Specification (SRS) document. The structure and content should be entirely based on the information discussed in the conversation. Follow these guidelines:

    1. Start with an introduction that summarizes the project.
    2. Create logical sections based on the topics discussed in the conversation.
    3. Include all relevant details mentioned, such as project goals, scope, features, requirements, constraints, and any other important aspects.
    4. Use appropriate headings and subheadings to organize the information.
    5. If certain standard SRS sections are applicable but not explicitly discussed, include them with a note that they require further discussion.
    6. Ensure the document flows logically and covers all aspects of the project mentioned in the conversation.

    Conversation History:
    {conversation_text}

    Generate the SRS document content:
    """

    srs_response = client.chat.completions.create(
        messages=[
            {"role": "system", "content": SYSTEM_MESSAGE_EN if user_language == 'en' else SYSTEM_MESSAGE_JP},
            {"role": "user", "content": srs_prompt}
        ],
        model="llama3-8b-8192",
    )
    return srs_response.choices[0].message.content

def create_srs_document(content):
    doc = Document()
    doc.add_heading('Software Requirements Specification (SRS)', 0)

    lines = content.split('\n')
    current_level = 0
    for line in lines:
        if line.strip():
            if line[0].isdigit() or line.isupper():
                level = len(line.split('.')) if '.' in line else (1 if line.isupper() else 2)
                doc.add_heading(line.strip(), level=level)
                current_level = level
            else:
                if line.startswith('  '):
                    doc.add_paragraph(line.strip(), style='List Bullet')
                else:
                    doc.add_paragraph(line.strip())

    return doc

# def process_response(response):
#     lines = response.split('\n')
#     lines = [line.strip() for line in lines]
#     lines = [line for line in lines if line]
#     processed = '\n'.join(lines)
#     processed = re.sub(r'\n{3,}', '\n\n', processed)
#     return processed


def process_response(response):
    # Split the response into paragraphs
    paragraphs = re.split(r'\n\s*\n', response.strip())
    
    # Process each paragraph
    processed_paragraphs = []
    for para in paragraphs:
        # Remove leading/trailing whitespace and join lines within a paragraph
        lines = [line.strip() for line in para.split('\n') if line.strip()]
        processed_para = ' '.join(lines)
        
        # Preserve Markdown list formatting
        processed_para = re.sub(r'(\d+\.\s|\-\s)', r'\n\1', processed_para)
        
        processed_paragraphs.append(processed_para)
    
    # Join paragraphs with a single newline
    processed = '\n\n'.join(processed_paragraphs)
    
    # Ensure code blocks are properly formatted
    processed = re.sub(r'```(\w+)\s*\n', r'```\1\n', processed)
    
    return processed

@app.route('/')
def home():
    return """
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>KUROCO LAB Chatbot</title>
        <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/5.15.3/css/all.min.css">
        <script src="https://cdnjs.cloudflare.com/ajax/libs/marked/2.0.3/marked.min.js"></script>
<style>
    :root {
        --primary-color: #6a11cb;
        --secondary-color: #2575fc;
        --accent-color: #4a00e0;
        --text-color: #2c3e50;
        --sidebar-width: 280px;
        --chat-bg: #ffffff;
        --user-msg-bg: #e6f3ff;
        --bot-msg-bg: #f0f0f0;
        --input-area-height: 100px;
        --input-bg: #f8f8f8;
    }

    body, html {
        font-family: 'Roboto', 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
        margin: 0;
        padding: 0;
        height: 100%;
        overflow: hidden;
    }

    .container {
        display: grid;
        grid-template-columns: var(--sidebar-width) 1fr;
        height: 100vh;
    }

    .sidebar {
        background: linear-gradient(135deg, #6a11cb 0%, #2575fc 100%);
        color: white;
        padding: 20px;
        overflow-y: auto;
        display: flex;
        flex-direction: column;
        justify-content: space-between;
        box-shadow: 2px 0 10px rgba(0, 0, 0, 0.1);
    }

    .sidebar h2 {
        margin-top: 0;
        font-size: 28px;
        text-align: center;
        text-transform: uppercase;
        letter-spacing: 2px;
        padding-bottom: 10px;
        border-bottom: 2px solid rgba(255, 255, 255, 0.3);
        text-shadow: 2px 2px 4px rgba(0, 0, 0, 0.3);
    }

    .sidebar ul {
        list-style-type: none;
        padding: 0;
    }

    .sidebar li {
        margin-bottom: 15px;
    }

    .sidebar a {
        color: white;
        text-decoration: none;
        font-size: 18px;
        transition: all 0.3s ease;
        display: block;
        padding: 10px;
        border-radius: 5px;
    }

    .sidebar a:hover {
        background-color: rgba(255, 255, 255, 0.2);
        transform: translateX(5px);
    }

    .sidebar-buttons {
        margin-top: 20px;
    }

    .sidebar-buttons button {
        background-color: rgba(255, 255, 255, 0.2);
        color: white;
        border: none;
        border-radius: 5px;
        cursor: pointer;
        transition: all 0.3s ease;
        font-size: 16px;
        font-weight: bold;
        padding: 12px;
        margin-bottom: 10px;
        width: 100%;
        text-shadow: 1px 1px 2px rgba(0, 0, 0, 0.2);
    }

    .sidebar-buttons button:hover {
        background-color: rgba(255, 255, 255, 0.3);
        transform: translateY(-2px);
        box-shadow: 0 4px 8px rgba(0, 0, 0, 0.2);
    }

    #language-selection {
        margin-top: 20px;
    }

    #language-selection h3 {
        font-size: 18px;
        margin-bottom: 10px;
        color: rgba(255, 255, 255, 0.8);
    }

    .language-option {
        display: flex;
        align-items: center;
        padding: 10px;
        cursor: pointer;
        transition: background-color 0.3s ease;
        border-radius: 5px;
        background-color: rgba(255, 255, 255, 0.1);
    }

    .language-option:hover {
        background-color: rgba(255, 255, 255, 0.2);
    }

    .language-option.active {
        background-color: rgba(255, 255, 255, 0.3);
    }

    .language-option img {
        width: 24px;
        height: 24px;
        margin-right: 10px;
        border-radius: 50%;
    }

    .language-label {
        font-size: 16px;
        color: white;
    }

    .copyright {
        font-size: 12px;
        text-align: center;
        margin-top: auto;
        padding-top: 20px;
        opacity: 0.7;
    }

    .main-content {
        display: flex;
        flex-direction: column;
        height: 100%;
        overflow: hidden;
        background-color: var(--chat-bg);
    }

    #chat-container {
        flex-grow: 1;
        display: flex;
        flex-direction: column;
        overflow: hidden;
        padding: 20px;
    }

    #chat-messages {
        flex-grow: 1;
        overflow-y: auto;
        display: flex;
        flex-direction: column-reverse;
        padding-bottom: 20px;
    }

    .message {
        margin-top: 20px;
        padding: 15px;
        border-radius: 8px;
        max-width: 80%;
        line-height: 1.5;
        opacity: 0;
        transform: translateY(20px);
        animation: slideIn 0.3s ease forwards;
    }

    @keyframes slideIn {
        to {
            opacity: 1;
            transform: translateY(0);
        }
    }

    .bot-message {
        background-color: var(--bot-msg-bg);
        color: var(--text-color);
        align-self: flex-start;
        border-left: 4px solid var(--primary-color);
        white-space: pre-wrap;
        font-family: 'Roboto', 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
        line-height: 1.6;
    }

    .user-message {
        background-color: var(--user-msg-bg);
        color: var(--text-color);
        align-self: flex-end;
        border-right: 4px solid var(--accent-color);
    }

    .bot-message p {
        margin-bottom: 0.8em;
    }

    .bot-message ul, .bot-message ol {
        margin-left: 1.5em;
        margin-bottom: 0.8em;
    }

    .bot-message li {
        margin-bottom: 0.4em;
    }

    .bot-message code {
        background-color: #e0e0e0;
        padding: 2px 4px;
        border-radius: 4px;
        font-family: 'Consolas', 'Monaco', 'Courier New', monospace;
    }

    .bot-message pre {
        background-color: #e0e0e0;
        padding: 10px;
        border-radius: 4px;
        overflow-x: auto;
        margin-bottom: 0.8em;
    }

    .bot-message pre code {
        background-color: transparent;
        padding: 0;
    }

    #typing-indicator {
        padding: 10px;
        background-color: var(--chat-bg);
        opacity: 0;
        transform: translateY(20px);
        transition: opacity 0.3s ease, transform 0.3s ease;
    }

    #typing-indicator.visible {
        opacity: 1;
        transform: translateY(0);
    }

    #user-input-container {
        display: flex;
        gap: 15px;
        align-items: center;
        background-color: var(--input-bg);
        padding: 15px;
        border-top: 1px solid #e0e0e0;
        height: var(--input-area-height);
    }

    #user-input {
        flex-grow: 1;
        padding: 15px;
        border: 1px solid #d0d0d0;
        border-radius: 8px;
        font-size: 16px;
        transition: all 0.3s ease;
        resize: none;
        height: 70px;
        background-color: #ffffff;
        color: var(--text-color);
    }

    #user-input::placeholder {
        color: #999999;
    }

    #user-input:focus {
        outline: none;
        border-color: var(--primary-color);
        box-shadow: 0 0 0 2px rgba(106, 17, 203, 0.2);
    }

    #send-button {
        padding: 15px 30px;
        background-color: var(--primary-color);
        color: white;
        border: none;
        border-radius: 8px;
        cursor: pointer;
        font-size: 18px;
        transition: all 0.3s ease;
    }

    #send-button:hover {
        background-color: var(--accent-color);
    }

    .loading-animation {
        display: flex;
        justify-content: center;
        align-items: center;
        margin-top: 20px;
    }

    .dot {
        width: 8px;
        height: 8px;
        background-color: var(--primary-color);
        border-radius: 50%;
        margin: 0 5px;
        animation: bounce 1.4s infinite ease-in-out both;
    }

    .dot:nth-child(1) { animation-delay: -0.32s; }
    .dot:nth-child(2) { animation-delay: -0.16s; }

    @keyframes bounce {
        0%, 80%, 100% { transform: scale(0); }
        40% { transform: scale(1); }
    }

    @media (max-width: 768px) {
        .container {
            grid-template-columns: 1fr;
        }

        .sidebar {
            position: fixed;
            left: -100%;
            top: 0;
            bottom: 0;
            width: 80%;
            max-width: 300px;
            z-index: 1000;
            transition: left 0.3s ease;
        }

        .sidebar.open {
            left: 0;
        }

        #menu-toggle {
            display: block;
            position: fixed;
            top: 10px;
            left: 10px;
            z-index: 1001;
            background: var(--primary-color);
            color: white;
            border: none;
            padding: 10px;
            font-size: 20px;
            cursor: pointer;
        }
    }
</style>
    </head>
    <body>
        <div class="container">
            <div class="sidebar">
                <div>
                    <h2>KUROCOLAB BOT</h2>
                    <ul>
                        <li><a href="#"><i class="fas fa-home"></i> <span class="menu-text" data-en="Home" data-jp="ホーム">Home</span></a></li>
                        <li><a href="#"><i class="fas fa-project-diagram"></i> <span class="menu-text" data-en="Projects" data-jp="プロジェクト">Projects</span></a></li>
                        <li><a href="#"><i class="fas fa-file-alt"></i> <span class="menu-text" data-en="Documents" data-jp="ドキュメント">Documents</span></a></li>
                        <li><a href="#"><i class="fas fa-cog"></i> <span class="menu-text" data-en="Settings" data-jp="設定">Settings</span></a></li>
                    </ul>
                </div>
                <div class="sidebar-buttons">
                    <button id="clear-chat"><i class="fas fa-trash"></i> <span class="button-text" data-en="Clear Chat" data-jp="チャットをクリア">Clear Chat</span></button>
                    <button id="export-chat"><i class="fas fa-download"></i> <span class="button-text" data-en="Export Chat" data-jp="チャットをエクスポート">Export Chat</span></button>
                </div>
                <div id="language-selection">
                    <h3 class="menu-text" data-en="Language" data-jp="言語">Language</h3>
                    <div class="language-option active" onclick="setLanguage('en')">
                        <img src="https://flagcdn.com/w40/gb.png" alt="English">
                        <span class="language-label">English</span>
                    </div>
                    <div class="language-option" onclick="setLanguage('jp')">
                        <img src="https://flagcdn.com/w40/jp.png" alt="日本語">
                        <span class="language-label">日本語</span>
                    </div>
                </div>
                <div class="copyright">
                    <h4>Ⓒ2024 FREECOMPANY Inc.</h4>
                </div>
            </div>
            <div class="main-content">
                <div id="chat-container">
                    <div id="chat-messages"></div>
                    <div id="typing-indicator" style="display: none;">
                        <div class="loading-animation">
                            <div class="dot"></div>
                            <div class="dot"></div>
                            <div class="dot"></div>
                        </div>
                    </div>
                </div>
                <div id="user-input-container">
                    <textarea id="user-input" placeholder="Type your message here..." data-en="Type your message here..." data-jp="メッセージを入力してください..."></textarea>
                    <button id="send-button"><i class="fas fa-paper-plane"></i></button>
                </div>
            </div>
        </div>

        <script src="https://cdnjs.cloudflare.com/ajax/libs/brython/3.9.0/brython.min.js"></script>
        <script src="https://cdnjs.cloudflare.com/ajax/libs/brython/3.9.0/brython_stdlib.js"></script>

        <script type="text/python">
        from browser import document, ajax, window
        import json

        user_language = 'en'

        def on_complete(req):
            response = json.loads(req.text)
            add_message(response['response'], 'bot')
            hide_typing_indicator()

        def send_message(event):
            user_input = document['user-input'].value
            if user_input.strip() == "":
                return
            add_message(user_input, 'user')
            document['user-input'].value = ""
            show_typing_indicator()
            
            req = ajax.Ajax()
            req.bind('complete', on_complete)
            req.open('POST', '/chat', True)
            req.set_header('content-type', 'application/json')
            req.send(json.dumps({'message': user_input, 'language': user_language}))

        def add_message(message, sender):
            chat_messages = document['chat-messages']
            new_message = document.createElement('div')
            new_message.classList.add('message', f'{sender}-message')
            if sender == 'bot':
                # Parse Markdown for bot messages
                new_message.innerHTML = window.marked(message)
            else:
                new_message.textContent = message
            chat_messages.insertBefore(new_message, chat_messages.firstChild)
            chat_messages.scrollTop = 0

        def show_typing_indicator():
            typing_indicator = document['typing-indicator']
            typing_indicator.style.display = 'block'
            window.setTimeout(lambda: typing_indicator.classList.add('visible'), 10)

        def hide_typing_indicator():
            typing_indicator = document['typing-indicator']
            typing_indicator.classList.remove('visible')
            window.setTimeout(lambda: setattr(typing_indicator.style, 'display', 'none'), 300)

        def clear_chat(event):
            document['chat-messages'].innerHTML = ''

        def export_chat(event):
            chat_content = document['chat-messages'].innerHTML
            
            req = ajax.Ajax()
            req.bind('complete', on_export_complete)
            req.open('POST', '/export-chat', True)
            req.set_header('content-type', 'application/json')
            req.send(json.dumps({'content': chat_content}))

        def on_export_complete(req):
            if req.status == 200:
                blob = window.Blob.new([req.text], {'type': 'text/html'})
                url = window.URL.createObjectURL(blob)
                a = document.createElement('a')
                a.href = url
                a.download = 'chat_export.html'
                a.click()
                window.URL.revokeObjectURL(url)
                
                content_size = len(req.text) / 1024  # size in KB
                window.alert(f"Chat exported successfully. Size: {content_size:.2f} KB")
            else:
                error_message = json.loads(req.text)['error']
                window.alert(f"Export failed: {error_message}")

        def set_language(lang):
            global user_language
            user_language = lang
            if lang == 'en':
                initial_message = "Hello! I'm the KUROCO LAB chatbot, your managing director for project implementation. How can I assist you today?"
                document.select_one('.language-option:nth-of-type(1)').classList.add('active')
                document.select_one('.language-option:nth-of-type(2)').classList.remove('active')
            else:
                initial_message = "こんにちは！KUROCO LABチャットボットです。プロジェクト実施のマネージングディレクターとして、本日はどのようなお手伝いができますか？"
                document.select_one('.language-option:nth-of-type(2)').classList.add('active')
                document.select_one('.language-option:nth-of-type(1)').classList.remove('active')
            document['chat-messages'].innerHTML = ''
            add_message(initial_message, 'bot')
            update_ui_text(lang)

        def update_ui_text(lang):
            elements = document.select('.menu-text, .button-text')
            for el in elements:
                el.text = el.attrs[f'data-{lang}']
            document['user-input'].attrs['placeholder'] = document['user-input'].attrs[f'data-{lang}']

        document['send-button'].bind('click', send_message)
        document['user-input'].bind('keypress', lambda event: send_message(event) if event.keyCode == 13 and not event.shiftKey else None)
        document['clear-chat'].bind('click', clear_chat)
        document['export-chat'].bind('click', export_chat)

        window.setLanguage = set_language

        # Initial bot message
        set_language('en')
        </script>

        <script type="text/javascript">
            window.onload = function() {
                brython();
            }
        </script>
    </body>
    </html>
    """

# @app.route('/chat', methods=['POST'])
# @limiter.limit("5 per minute")
# def chat():
#     global user_language
#     try:
#         user_message = request.json['message']
#         user_language = request.json['language']
#         if not user_message or not isinstance(user_message, str):
#             raise BadRequest("Invalid message format")
        
#         conversation_history.append(user_message)
        
#         system_message = SYSTEM_MESSAGE_EN if user_language == 'en' else SYSTEM_MESSAGE_JP
        
#         response = client.chat.completions.create(
#             messages=[
#                 {"role": "system", "content": system_message},
#                 {"role": "system", "content": "Format your responses concisely, using Markdown. Use a single newline between paragraphs, and avoid excessive spacing. Use **bold** for emphasis, - for unordered lists, 1. for ordered lists, and `code` for inline code or ```language for code blocks."},
#                 *[{"role": "user" if i % 2 == 0 else "assistant", "content": msg} for i, msg in enumerate(conversation_history)]
#             ],
#             model="llama3-8b-8192",
#         )
        
#         response_content = response.choices[0].message.content
#         processed_response = process_response(response_content)
#         processed_response = process_assistant_message(processed_response, user_message)
#         conversation_history.append(processed_response)
        
#         return jsonify({'response': processed_response})
#     except Exception as e:
#         app.logger.error(f"An error occurred: {str(e)}")
#         raise InternalServerError("An unexpected error occurred")



# Update the chat function to use the new process_response
@app.route('/chat', methods=['POST'])
@limiter.limit("5 per minute")
def chat():
    global user_language
    try:
        user_message = request.json['message']
        user_language = request.json['language']
        if not user_message or not isinstance(user_message, str):
            raise BadRequest("Invalid message format")
        
        conversation_history.append(user_message)
        
        system_message = SYSTEM_MESSAGE_EN if user_language == 'en' else SYSTEM_MESSAGE_JP
        
        response = client.chat.completions.create(
            messages=[
                {"role": "system", "content": system_message},
                {"role": "system", "content": "Format your responses concisely, using Markdown. Use a single newline between paragraphs. Use **bold** for emphasis, - for unordered lists, 1. for ordered lists, and `code` for inline code or ```language for code blocks. Avoid unnecessary spacing."},
                *[{"role": "user" if i % 2 == 0 else "assistant", "content": msg} for i, msg in enumerate(conversation_history)]
            ],
            model="llama3-8b-8192",
        )
        
        response_content = response.choices[0].message.content
        processed_response = process_response(response_content)
        processed_response = process_assistant_message(processed_response, user_message)
        conversation_history.append(processed_response)
        
        return jsonify({'response': processed_response})
    except Exception as e:
        app.logger.error(f"An error occurred: {str(e)}")
        raise InternalServerError("An unexpected error occurred")

@app.route("/create_document/<doc_id>", methods=["GET"])
def get_document(doc_id):
    try:
        if doc_id not in documents:
            raise NotFound("Document not found")
        content = documents[doc_id]
        doc = create_srs_document(content)
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return send_file(
            doc_io,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='SRS_Document.docx'
        )
    except Exception as e:
        app.logger.error(f"An error occurred while creating the document: {e}")
        raise InternalServerError("Failed to create document")

@app.route('/export-chat', methods=['POST'])
def export_chat():
    try:
        chat_content = request.json['content']
        content_bytes = chat_content.encode('utf-8')
        
        # Check size (limit to 10 MB)
        max_size = 10 * 1024 * 1024  # 10 MB in bytes
        if len(content_bytes) > max_size:
            return jsonify({
                'error': f"Chat export is too large ({len(content_bytes) / 1024 / 1024:.2f} MB). Maximum size is {max_size / 1024 / 1024} MB."
            }), 413  # 413 Payload Too Large
        
        response = make_response(chat_content)
        response.headers.set('Content-Type', 'text/html')
        response.headers.set('Content-Disposition', 'attachment', filename='chat_export.html')
        return response

    except Exception as e:
        app.logger.error(f"An error occurred during chat export: {str(e)}")
        raise InternalServerError("An unexpected error occurred during chat export")

@app.errorhandler(BadRequest)
@app.errorhandler(NotFound)
@app.errorhandler(InternalServerError)
def handle_error(error):
    return jsonify({'error': str(error)}), error.code

if __name__ == '__main__':
    app.run(debug=True) 